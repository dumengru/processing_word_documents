#!/usr/bin/python
# coding=UTF-8

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import  Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import  qn
from docx.shared import Inches
from docx.shared import Cm
#打开文档
document = Document()

#报告标题
head = document.add_heading()
head.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
head.paragraph_format.space_before = Pt(36)
run = head.add_run(u"生成的报告")
run.font.size=Pt(36)
run.font.color.rgb=RGBColor(0x00,0x00,0x00) #标题颜色
document.styles['Normal'].font.name = u'Arial' #字体1
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') #汉字使用字体2->宋体
#添加分页
document.add_page_break()

#二级标题
head2 = document.add_heading(level=2)
head2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = head2.add_run(u'报告分析')
#run.font.name=u'宋体'
run.font.size=Pt(21)
run.font.color.rgb=RGBColor(14,106,85)
#段落后行距
head2.paragraph_format.space_after = Pt(30)

#二级级标题段落
#document.styles['Normal'].font.name = u'宋体'
p = document.add_paragraph()
run = p.add_run(u'''    段落1
	段落2''')
run.font.color.rgb = RGBColor(0, 0, 0)
run.font.size = Pt(15)
#单倍行距
p.paragraph_format.line_spacing = Pt(30)
#段落后行距
p.paragraph_format.space_after = Pt(30)

#三级标题
head3 = document.add_heading(level=3)
run = head3.add_run(u'分析对象:')
#run.font.name=u'宋体'
run.font.size=Pt(18)
run.font.color.rgb=RGBColor(0,0,0)
#段落后行距
head2.paragraph_format.space_after = Pt(30)
#基本信息
document.add_picture('./image/1.png', width=Cm(15.24)) #设置图片大小，这里用的是cm厘米，比较直观吧，反正我从来不懂Inches跟cm之间怎么换，也不知他有多大

#设置中文字体
#run = paragraph.add_run(u'设置中文字体，')
#run.font.name=u'宋体'
#r = run._element
#r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#设置粗体小四
paragraph = document.add_paragraph()
run = paragraph.add_run(u'MD5:  ')
run.bold = True
run.font.size=Pt(12)
paragraph = document.add_paragraph()
run = paragraph.add_run(u'TYPE:  ')
run.bold = True
run.font.size=Pt(12)
paragraph = document.add_paragraph()
run = paragraph.add_run(u'文件名称:  ')
run.bold = True
run.font.size=Pt(12)
paragraph = document.add_paragraph()
run = paragraph.add_run(u'文件大小:  ')
run.bold = True #加粗
run.font.size=Pt(12) #小四

#报告1
document.add_picture('./image/2.png', width=Cm(15.24))
paragraph = document.add_paragraph()
run = paragraph.add_run(u'判定:  ')
run.bold = True
run.font.size=Pt(12)

#报告2
document.add_picture('./image/3.png', width=Cm(15.24))
paragraph = document.add_paragraph()
run = paragraph.add_run(u'报告2:  ')
run.bold = True
run.font.size=Pt(12)

#分析总结
document.add_picture('./image/4.png', width=Cm(15.24))
paragraph = document.add_paragraph()
run = paragraph.add_run(u'分析结果:  ')
run.bold = True
run.font.size=Pt(12)

#检测结果
document.add_picture('./image/5.png', width=Cm(15.24))
paragraph = document.add_paragraph()
run = paragraph.add_run(u'检测结果:  ')
run.bold = True
run.font.size=Pt(12)

#关系
document.add_picture('./image/6.png', width=Cm(15.24))
paragraph = document.add_paragraph()
run = paragraph.add_run(u'关系:  ')
run.bold = True #加粗
run.font.size=Pt(12)#小四

#关键字符串
document.add_picture('./image/7.png', width=Cm(15.24))
paragraph = document.add_paragraph()
run = paragraph.add_run(u'关键字符串:  ')
run.bold = True
run.font.size=Pt(12)

head4 = document.add_heading(level=3)
run = head4.add_run(u'附件:')
run.font.size=Pt(18)
run.font.color.rgb=RGBColor(0,0,0)

#注1
head5 = document.add_heading(level=3)
run = head5.add_run(u'注1:')
run.font.size=Pt(15)
run.font.color.rgb=RGBColor(65,186,126)
head5.paragraph_format.line_spacing = Pt(30)

p = document.add_paragraph()
p.paragraph_format.line_spacing = Pt(30)
run = p.add_run(u'''    注解1段落''')
run.font.color.rgb = RGBColor(0, 0, 0)
run.font.size = Pt(12)

#注2
head6 = document.add_heading(level=3)
run = head6.add_run(u'注2:')
run.font.size=Pt(15)
run.font.color.rgb=RGBColor(65,186,126)
head6.paragraph_format.line_spacing = Pt(30)

p = document.add_paragraph()
p.paragraph_format.line_spacing = Pt(30)
run = p.add_run(u'''    注解2段落''')
run.font.color.rgb = RGBColor(0, 0, 0)
run.font.size = Pt(12)

#保存文件
document.save('demo.docx')
