import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT           # 设置对象居中, 对齐等
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER  # 设置制表符
from docx.enum.table import WD_TABLE_ALIGNMENT             # 设置表格对齐方式
from docx.enum.text import WD_LINE_SPACING                  # 设置行间距
from docx.shared import Inches                              # 设置图像大小: 英寸
from docx.shared import Pt                                  # 设置字号: 磅值
from docx.shared import RGBColor                            # 设置字体颜色
from docx.shared import Length                              # 设置宽度
from docx.shared import Cm                                  # 厘米单位
from docx.oxml.ns import qn                                 # 中文字体

m_docx = docx.Document()    # 可以直接打开 "test_read.docx"


def write_title():
    """
    1. 添加标题
    2. 添加正文
    3. 添加分页符

    保存文件
    """
    m_docx.add_heading("添加一级标题", level=1)            # 添加一级标题
    m_docx.add_paragraph("添加正文")                      # 添加正文

    block = m_docx.add_paragraph("后文添加块(设置格式):\n")      # 开始添加块内容
    block.add_run("加粗\n").bold = True                        # 添加块格式
    block.add_run("普通\n")
    block.add_run("斜体\n").italic = True                      # 添加斜体
    block.add_run("下划线\n").underline = True                 # 添加下划线
    block.add_run("删除线\n").strike = True                    # 添加删除线
    block.add_run("阴影\n").shadow = True                     # 添加阴影
    block.add_run("24号字\n").size = Pt(24)                   # 24号字
    block.add_run("添加颜色\n").rgb = RGBColor(255, 0, 0)     # 添加颜色
    block.add_run("English\n").name = "Arial"               # 英文字体

    para = m_docx.paragraphs[1]                         # 开始添加段落: 获取第二段
    para.insert_paragraph_before("新增段落: 第二段")      # 第二段之前添加

    m_docx.add_page_break()                             # 添加分页符
    m_docx.add_paragraph("下一页内容")

    m_docx.add_picture("word_pic/test_word_image1.png", # 添加图片， 设置宽高
                       width=Cm(13), height=Cm(8))

    m_docx.add_table(rows=3, cols=3, style="Medium Grid 1 Accent 1")  # 插入单元格, 添加预设样式

    # 在表格中添加图片步骤
    # 1. 定位表格
    # 2. 在表格中添加块
    #
    run = m_docx.tables[0].cell(1, 1).paragraphs[0].add_run()
    run.add_picture("word_pic/test_word_image2.png",    # 添加图片， 设置宽高
                       width=Cm(3), height=Cm(3))

    # 查询文件中图片总数量
    print("图片总数量: ", len(m_docx.inline_shapes))
    print("图表宽高: ", m_docx.inline_shapes[0].width, "_", m_docx.inline_shapes[0].height)

    # 添加图片(文字)右对齐
    m_docx.paragraphs[-1].add_run().add_picture("word_pic/test_word_image3.jpeg", width=Cm(3), height=Cm(3))
    m_docx.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # 表格增加一行和一列
    m_docx.tables[0].add_row()
    m_docx.tables[0].add_column(width=Cm(2))    # 增加列必须设置列宽

    # 查看表格行列数
    print("表格行列: ", len(m_docx.tables[0].rows), "_", len(m_docx.tables[0].columns))

    # 获取表第二行, 删除表第二行
    row2 = m_docx.tables[0].rows[1]
    row2._element.getparent().remove(row2._element)
    print("删除表格第二行")
    print("表格行列: ", len(m_docx.tables[0].rows), "_", len(m_docx.tables[0].columns))

    # 向表格添加数据: 两种方式
    # 1. 直接赋值
    # 2. 获取表中的段落
    table0 = m_docx.tables[0]
    table0.cell(0, 0).text = "表格添加内容"

    table_para = m_docx.tables[0].cell(0, 1).paragraphs[0]
    table_para.text = "表格段落"
    table_add_para = m_docx.tables[0].cell(0, 2).add_paragraph("表格添加段落")
    table_add_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    m_docx.add_paragraph("添加段落English 1\n")
    m_docx.add_paragraph("添加段落English 2\n")
    m_docx.add_paragraph("添加段落English 3\n")
    for para in m_docx.paragraphs[-3:]:

        for run in para.runs:
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 设置全局正文
    m_docx.styles["Normal"].font.name = "Arial"
    m_docx.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    # 设置标题:
    head = m_docx.add_heading("", level=1)
    head.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    head.paragraph_format.space_before = Pt(36)
    run = head.add_run(u"添加标题")
    run.font.size = Pt(12)
    run.font.name = u"黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 标题颜色

    # 获取字体大小
    # 将正文居中对齐, 段落行间距, 一级标题右对齐
    for para in m_docx.paragraphs:
        print(para.style.name)
        if para.style.name == "Normal":
            # para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # para.paragraph_format.line_spacing = 3.0
            # 1.5倍行间距: ONE_POINT_FIVE
            # 最小行距: AT_LEAST
            # 双倍行间距: DOUBLE,
            # 固定值: EXACTLY,(需设置字号)
            # 多倍行间距: MULTIPLE,
            # SINGL, 单倍行间距
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = Pt(18)

            # 段前段后
            para.paragraph_format.space_before = Pt(18)
            para.paragraph_format.space_after = Pt(12)

            # 左缩进, 首行缩进, 悬挂缩进
            para.paragraph_format.left_indent = Pt(0.3)
            para.paragraph_format.first_line_indent = Inches(1)


        if para.style.name == "Heading 1":
            para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT



    m_docx.save("test_write.docx")


# def add_pic():

# 定位段落
def location_para():
    m_docx = docx.Document("test_read.docx")
    print(len(m_docx.paragraphs))


if __name__ == '__main__':
    write_title()