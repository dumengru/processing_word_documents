import docx
from docx import Document   #用来建立一个word对象
from docx.shared import Pt  #用来设置字体的大小
from docx.shared import Inches
from docx.oxml.ns import qn  #设置字体
from docx.shared import RGBColor  #设置字体的颜色
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  #设置对其方式
from docx.shared import Cm
from docx.enum import style


#创建一个空白的word文档
doc=Document()
doc.styles["Normal"].font.name=u"Times New Roman"  #设置全局字体
doc.styles["Normal"]._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')#设置中文字体
doc.styles["Normal"].font.color.rgb=RGBColor(255,0,0)#设置正文全局颜色为红色
doc.styles["Normal"].font.size=Pt(29)#设置正文全局大小为29
doc.styles["Heading 2"].font.size=Pt(29)#设置全局2级标题的字体大小为29



# 获取段落样式
paragraph_format = style.paragraph_format
# 首行缩进0.74厘米，即2个字符
paragraph_format.first_line_indent = Cm(0.74)




p=doc.add_paragraph()
r=p.add_run("违反接收")#这个段落使用的是Normal的样式

para_heading=doc.add_heading('', level=2)#返回1级标题段落对象，标题也相当于一个段落
run=para_heading.add_run(u"前言")#这个使用的是"Heading 2" 的样式

doc.save("特殊字体.docx")


# 增加分页
doc.add_page_break()

par2 = doc.add_paragraph('第二个段落文本:麻烦给我的爱人来一杯MojIto，'
                         '我喜欢阅读她微醺时的眼眸，而我的咖啡，糖不用太多,'
                         '这世界已经因为她甜得过头，没有跟她笑容一样浓郁的雪茄，'
                         '就别浪费时间介绍收起来吧，拱廊的壁画，旧城的涂鸦，'
                         '所有色彩都因为她说不出话。')

# Paragraph 段落格式设置
# 段落对齐设置：# 导入对齐选项
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 设置段落文本右对齐
par2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
'''
可选项有：
'LEFT', 0, 'left', 'Left-aligned'
'CENTER', 1, 'center', 'Center-aligned.'
'RIGHT', 2, 'right', 'Right-aligned.'
'JUSTIFY', 3, 'both', 'Fully justified.'
'DISTRIBUTE', 4, 'distribute', 'Paragraph characters are distrib'
            'uted to fill the entire width of the paragraph.'
'JUSTIFY_MED', 5, 'mediumKashida', 'Justified with a medium char'
            'acter compression ratio.'
'JUSTIFY_HI', 7, 'highKashida', 'Justified with a high character'
            ' compression ratio.'
'JUSTIFY_LOW', 8, 'lowKashida', 'Justified with a low character '
            'compression ratio.'
'THAI_JUSTIFY', 9, 'thaiDistribute', 'Justified according to Tha'
            'i formatting layout.'
'''
# 获取段落设置的对齐方式
print(par2.alignment)





# 段落缩进：
# 导入缩进单位
from docx.shared import Inches, Pt
# 左缩进,0.5 英寸
par2.paragraph_format.left_indent = Inches(0.5)
# 右缩进,20 磅
par2.paragraph_format.right_indent = Pt(20)
# 首行缩进
par2.paragraph_format.first_line_indent = Inches(1)


# 行间距
# par2.paragraph_format.line_spacing = Pt(20) # 20 磅
par2.paragraph_format.line_spacing = 1.5 # 1.5倍行距 单倍行距 1.0
"当line_spacing设置为长度值时表示绝对距离，"
"设置为浮点数时表示行高的倍数"

# 段前间距
par2.paragraph_format.space_before = Pt(5)
# 段后间距
par2.paragraph_format.space_after = Pt(10)

# 设置段落内部文字在遇到需分页情况时处理状态
par2.paragraph_format.keep_together = True  # 段中不分页
par2.paragraph_format.keep_with_next = True # 与下段同页
par2.paragraph_format.page_break_before = True # 段前分页
par2.paragraph_format.widow_control = True # 孤行控制


# 获取段落的左缩进，首行缩进，段前间距：
l_space = par2.paragraph_format.left_indent
h_space = par2.paragraph_format.first_line_indent
b_space = par2.paragraph_format.space_before
print(l_space,h_space,b_space)
# 打印结果
# 457200 914400 63500


# 获取上面段落对象 所有 run 对象
par2_run = par2.runs
print(par2_run)
# 打印结果：
# [<docx.text.run.Run object at 0x000000000B307048>]


# 新建一个段落，增加一段文字
run_2 = doc.add_paragraph().add_run('这爱不落幕，忘了心事的国度，'
                                    '你所在之处，孤单都被征服，'
                                    '铁铸的招牌，错落着就像，'
                                    '一封封城市献给天空的情书，'
                                    '当街灯亮起，Havana漫步，'
                                    '这是世上最美丽的那双人舞。')
# 设置字体格式
run_2.font.name = 'Times New Roman'  # 注：这个好像设置 run 中的西文字体
# 设置中文字体
# 需导入 qn 模块
from docx.oxml.ns import qn
run_2.font.element.rPr.rFonts.set(qn('w:eastAsia'),'楷体')
# 设置字体大小
run_2.font.size = Pt(14)
# 设置加粗
run_2.font.bold = True
# 设置斜体
run_2.font.italic = True
# 设置字体颜色 需导入 rgb 颜色模块
from docx.shared import RGBColor
run_2.font.color.rgb = RGBColor(255,55,55) # 红色
# 设置下划线
# run_2.font.underline = True
# 设置轮廓线
# run_2.font.outline = True
# 设置阴影
# run_2.font.shadow = True
# 删除线
# run_2.font.strike = True
# 双删除线
# run_2.font.double_strike = True
# 设置下标
# run_2.font.subscript = True
# 设置上标
# run_2.font.superscript = True


#添加标题
def AddHeadText(text, size):
    title_ = document.add_heading(level=3)
    # title_.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER# 标题居中
    title_run = title_.add_run(text)  # 添加标题内容
    title_run.font.size = Pt(size)  # 设置标题字体大小
    title_run.font.name = 'Times New Roman'  # 设置标题西文字体
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')  # 设置标题中文字体
    title_run.font.color.rgb = RGBColor(0, 0, 0)#字体颜色


#添加段落内容(参数1：文本内容，参数2：字体大小，参数3：上行距,参数4：字体粗细，参数5：段落位置)
def AddParaText(text, size, space, thickness, position):
    p = document.add_paragraph()  # 段落
    #判断居中还是靠左,0为靠左
    if position == 0:
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT #靠左
    else :
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #居中
    p.paragraph_format.space_before = Pt(space)
    text = p.add_run(text)
    #判断字体是否加粗（1为不加粗）
    if thickness == 1:
        text.bold = False
    else:
        text.bold = True #加粗
    text.font.name = 'Times New Roman'
    text.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    text.font.size = Pt(size)



'''
遇到问题没人解答？小编创建了一个Python学习交流QQ群：857662006 
寻找有志同道合的小伙伴，互帮互助,群里还有不错的视频学习教程和PDF电子书！
'''
document=Document()
document.styles['Normal'].font.name=u'微软雅黑'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')


run = document.add_heading('', level=3).add_run(u"应用场景示例: ")#应用场景示例标题
run.font.name=u'微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')



Doc = Document()
Doc.styles['Normal'].font.name = u'宋体'
Doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
Doc.styles['Normal'].font.size = Pt(10.5)
Doc.styles['Normal'].font.color.rgb = RGBColor(0,0,0)
Head = Doc.add_heading("",level=1)# 这里不填标题内容
run  = Head.add_run("刚来csdn，这就是博客么，I了")
run.font.name=u'Cambria'
run.font.color.rgb = RGBColor(0,0,0)
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
Doc.add_paragraph("Python ")
Doc.add_paragraph("Python 对word进行操作")
Doc.save("Python_word.docx")



