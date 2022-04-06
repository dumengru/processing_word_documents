from docx_setting import WORD_SETTING, ALIGNMENT_DICT, HEADING_NUMBER
from docx import Document
from docx.shared import (
    Pt,                                 # 磅值
    Cm,                                 # 厘米单位
    Length,
)
from docx.oxml.ns import qn                                 # 中文字体
import re
from os import remove, listdir


class DocxModel():
    def __init__(self, file_name: str):
        """ 传入将要修正的文件路径 """
        self.file_name = file_name
        self.old_docx = Document(self.file_name)
        self.new_docx = Document()

    def init_format(self):
        """
        按照配置设置文件格式
         1. 读取文件
         2. 设置全部正文
         3. 设置标题: 通过正则表达式定位
         4. 设置段落
         5. 设置表格和图片
        """
        # 新增内容
        # 1. 设置正文
        self.new_docx.styles["Normal"].font.name = WORD_SETTING["英文字体"]
        self.new_docx.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_SETTING["中文字体"])
        self.new_docx.styles["Normal"].font.size = Pt(WORD_SETTING["正文字号"])
        self.new_docx.styles["Normal"].font.bold = WORD_SETTING["正文加粗"]


        # 按照顺序读取并添加段落,图片和表格
        # 1. 计算表格和图片数量
        total_table = len(self.old_docx.tables)
        total_picture = len(self.old_docx.inline_shapes)
        print("表格数量: ", total_table, "\t图片数量: ", total_picture)
        self.save_old_picture()         # 将旧文件保存在本地

        pattern_heading = re.compile(r"^\$(.级标题)\$(.*?)$", re.M | re.S)
        pattern_list = re.compile(r"^\$(列表)\$(.*?)$", re.M | re.S)
        pattern_table = re.compile(r"^\$表格\$$", re.M | re.S)
        pattern_picture = re.compile(r"^\$图片\$$", re.M | re.S)

        count_table = 0         # 表格计数
        count_picture = 0       # 图片计数

        for para in self.old_docx.paragraphs:                               # 遍历段落
            if len(para.text) == 0:     # 过滤空行
                continue
            match_heading = pattern_heading.match(para.text)
            match_table = pattern_table.match(para.text)
            match_picture = pattern_picture.match(para.text)
            match_list = pattern_list.match(para.text)
            if match_heading:
                level = match_heading.group(1)                              # 获取标题内容
                m_text = match_heading.group(2)
                m_level = HEADING_NUMBER[level]
                self.new_docx.add_heading(m_text, level=m_level)            # 设置标题级别
            elif match_list:
                m_text = match_list.group(2)                             # 获取列表内容
                new_para = self.new_docx.add_paragraph(text=m_text, style="List Bullet")
            elif match_table:
                old_table = self.old_docx.tables[count_table]
                new_table = self.new_docx.add_table(rows=len(old_table.rows), cols=len(old_table.columns), style=WORD_SETTING["表格样式"])
                for row_index, row in enumerate(old_table.rows):
                    for col_index, cell in enumerate(row.cells):
                        new_table.cell(row_index, col_index).text = old_table.cell(row_index, col_index).text

                count_table += 1
            elif match_picture:
                # 查询图片名 -> 添加段落,添加块,添加图片,设置图片格式 -> 删除图片
                count_picture += 1
                pic_name = f"word_image{count_picture}"
                pic_path = [i for i in listdir(".") if pic_name in i][0]
                new_para = self.new_docx.add_paragraph()
                new_para.add_run().add_picture(pic_path, width=Cm(WORD_SETTING["图片宽度"]), height=Cm(WORD_SETTING["图片高度"]))
                new_para.alignment = ALIGNMENT_DICT[WORD_SETTING["图片对齐"]]
                # 段前/后间距
                new_para.paragraph_format.space_before = 2
                new_para.paragraph_format.space_after = 2
                # 首行缩进
                new_para.paragraph_format.first_line_indent = Pt(0)
                remove(pic_path)
            else:
                new_para = self.new_docx.add_paragraph(text=para.text)
                # 段前/后间距
                new_para.paragraph_format.space_before = 1
                new_para.paragraph_format.space_after = 1
                # 首行缩进
                new_para.paragraph_format.first_line_indent = Pt(Length(new_para.style.font.size).pt * 2.0)

    def save_old_picture(self):
        """ 保存旧文件全部图片 """
        dict_rel = self.old_docx.part._rels
        for rel in dict_rel:
            rel = dict_rel[rel]
            if "image" in rel.target_ref:
                image_name = rel.target_ref.split("/")[-1]
                with open(f"word_{image_name}", "wb") as f:
                    f.write(rel.target_part.blob)

    def add_number(self):
        """ 为不同级别标题添加不同抬头, 首行缩进 """
        head1, head2, head3, head4, head5, head6, head7 = 0, 0, 0, 0, 0, 0, 0
        for para in self.new_docx.paragraphs:
            if para.style.name == "Normal":
                # 左缩进, 首行缩进, 悬挂缩进
                para.paragraph_format.left_indent = Pt(0)
                # para.paragraph_format.first_line_indent = Pt(Length(para.style.font.size).pt * 2.3)

            elif para.style.name == "Heading 1":
                head1 += 1
                para.text = f"第{head1}章 " + para.text
                head2, head3, head4, head5, head6, head7 = 0, 0, 0, 0, 0, 0

            elif para.style.name == "Heading 2":
                head2 += 1
                para.text = f"{head1}.{head2} " + para.text
                head3, head4, head5, head6, head7 = 0, 0, 0, 0, 0

            elif para.style.name == "Heading 3":
                head3 += 1
                para.text = f"{head1}.{head2}.{head3} " + para.text
                head4, head5, head6, head7 = 0, 0, 0, 0

            elif para.style.name == "Heading 4":
                head4 += 1
                para.text = f"{head1}.{head2}.{head3}.{head4} " + para.text
                head5, head6, head7 = 0, 0, 0

            elif para.style.name == "Heading 5":
                head5 += 1
                para.text = f"{head1}.{head2}.{head3}.{head4}.{head5} " + para.text
                head6, head7 = 0, 0

            elif para.style.name == "Heading 6":
                head6 += 1
                para.text = f"{head1}.{head2}.{head3}.{head4}.{head5}.{head6} " + para.text
                head7 = 0

            elif para.style.name == "Heading 7":
                head7 += 1
                para.text = f"{head1}.{head2}.{head3}.{head4}.{head5}.{head6}.{head7} " + para.text

    def set_format(self):
        """
        设置文本格式
        """
        # 最后调整段落格式
        for para in self.new_docx.paragraphs:
            if para.style.name == "Normal":
                para.paragraph_format.space_before = 1
                para.paragraph_format.space_after = 1
                # para.paragraph_format.line_spacing = 0            # 段间
                para.paragraph_format.first_line_indent = Cm(0.88)     # 首行缩进
            elif para.style.name == "Heading 1":                    # 标题居中对齐
                para.style.paragraph_format.alignment = ALIGNMENT_DICT[0]
            # 设置标题字体
            elif para.style.name.startswith("Heading"):
                for run in para.runs:
                    run.font.name = WORD_SETTING["英文字体"]
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_SETTING["中文字体"])

    def run(self):
        self.init_format()   # 调整格式
        self.add_number()   # 修改标题样式
        self.set_format()   # 
        self.new_docx.save(f"标准文档_{self.file_name}")


def word_tip():
    """ 程序提示 """
    print("该程序适合多人合作word格式修改")
    print(
        "kingstar_word文档模板使用说明: \n",
        "1. 该程序只能处理 docx 文档\n",
        "2. word内容全部按正文书写\n",
        "3. 一级标题文字前添加双$$标识 $一级标题$\n",
        "4. 图片和表格上方添加双$$标识 $图片$ $表格$",
        "5. 列表前添加双$$标识 $列表$"
    )
    print(WORD_SETTING)
    print(
        "\n程序运行提示:\n",
        "1. 将源文件和程序放在同一目录下\n"
        "2. 输出文件名为 【out_源文件名】\n",
    )


if __name__ == '__main__':
    word_tip()  # 程序提示
    f_input = input("请输入要处理的文件名(无后缀): ")
    file_name = f"{f_input}.docx"
    m_docx = DocxModel(file_name)
    m_docx.run()
