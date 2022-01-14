import docx


m_doc = docx.Document("test_read.docx")


# 提取表格
def get_table():
    flag = 0
    for table in m_doc.tables:
        for row_index, row in enumerate(table.rows):    # 按行遍历
            for col_index, cell in enumerate(row.cells):    # 获取行中单元格
                if "量化" in cell.text:
                    flag += 1

    print("量化一词在表中出现: ", flag)


# 提取图片
def get_pic():
    dict_rel = m_doc.part._rels
    for rel in dict_rel:
        rel = dict_rel[rel]
        print(rel.target_ref)
        if "image" in rel.target_ref:
            image_name = rel.target_ref.split("/")[-1]
            word_name = "test_word"
            save_path = "word_pic"
            with open("{}/{}".format(save_path, word_name+"_"+image_name), "wb") as f:
                f.write(rel.target_part.blob)


# 获取段落
def get_para():
    m_para = m_doc.paragraphs
    for para in m_para:
        print(para.text)


# 获取块
def get_run():
    """
    1. 获取段落
    2. 获取段落中的块: runs
    """
    for para in m_doc.paragraphs[:2]:
        for run in para.runs:
            print(run.text)


# 文字计数(不包含表格)
def get_count():
    count = 0
    for para in m_doc.paragraphs:
        for run in para.runs:
            if "量化" in run.text:
                count += 1
    print("量化一词在文档中出现: ", count)


# 获取所有标题(重要)
from docx.enum.style import WD_STYLE_TYPE

def get_title():
    """
    """
    title = m_doc.styles
    for i in title:
        if i.type == WD_STYLE_TYPE.PARAGRAPH:
            print(i.name)


# 获取标题内容
def get_para_style():
    for para in m_doc.paragraphs:
        # if para.style.name == "Heading 1":   # 如果段落等于标题一
        #     print("标题一: ", para.text)
        # if para.style.name == "Heading 2":  # 如果段落等于标题二
        #     print("标题二: ", para.text)
        if para.style.name == "Title":  # 如果段落等于标题六
            print("标题六: ", para.text)


# 获取所有标题( 正则表达式 )
import re
def get_all_title():
    for para in m_doc.paragraphs:
        if re.match("^Heading \d+$", para.style.name):  # 获取所有段落
            print(para.text)


# 获取所有正文
def get_text_style():
    for para in m_doc.paragraphs[:50]:
        if para.style.name == "Normal":   # 如果段落等于标题一
            print("正文: ", para.text)


if __name__ == '__main__':
    get_pic()
    # get_para_style()
